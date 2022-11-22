import docx


# Defince test data
testData = {
    'name': 'John Doe',
    'age': 42,
    'address': '123 Main Street',
    'city': 'Anytown'
}

# define test data for table
testDataRows = [
        {
            'name': 'John Doe',
            'age': 42,
            'address': '123 Main Street'
        },
        {
            'name': 'Martin Gomez',
            'age': 32,
            'address': '456 Main Street'
        },
        {
            'name': 'Peter Parker',
            'age': 36,
            'address': '456 Main Street'
        },
        {
            'name': 'Leon Kennedy',
            'age': 29,
            'address': '456 Main Street'
        },
    ]

class DocxTemplate:
    def __init__(self, templatePath):
        self.templatePath = templatePath
        self.doc = docx.Document(templatePath)

    def fillTemplate(self, data):
        for paragraph in self.doc.paragraphs:
            inline = paragraph.runs
            for i in range(len(inline)):
                inline[i].text = self.replaceText(inline[i].text, data)

    def getTables(self):
        ret = []
        for table in self.doc.tables:
            if len(table.rows) >= 2:
                ret.append({'type': table.rows[2].cells[0].text, 'object': table})
        return ret

    def fillTable(self, table, data):
        # save cell text to list
        cellText = []
        for cell in table.rows[1].cells:
            cellText.append(cell.text)

        # Remove first row (template row)
        table._tbl.remove(table.rows[2]._tr)
        table._tbl.remove(table.rows[1]._tr)
        
        # Add new rows
        for dataRow in data:
            table.add_row()
            # Add data to cells
            for i in range(len(cellText)):
                table.rows[-1].cells[i].text = self.replaceText(cellText[i], dataRow)

    def replaceText(self, templateText, data):
        for key in data:
            if '${' + key + '}' in templateText:
                templateText = templateText.replace('${' + key + '}', str(data[key]))
        return templateText

    def save(self, path):
        self.doc.save(path)

docTemplate = DocxTemplate('template.docx')
tables = docTemplate.getTables()
for table in tables:
    print('found table: ' + table['type'])
    docTemplate.fillTable(table['object'], testDataRows)
docTemplate.fillTemplate(testData)
docTemplate.save('outputDoc.docx')